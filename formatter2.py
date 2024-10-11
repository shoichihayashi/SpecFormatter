from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement  # for working with page numbers
from docx.oxml.ns import qn  # Import for XML namespaces

# Load source and target documents
source_doc = Document('source.docx')
target_doc = Document('target.docx')

# Extract formatting from a paragraph run
def extract_formatting(run):
    formatting = {}
    formatting['bold'] = run.bold if run.bold is not None else False
    formatting['italic'] = run.italic if run.italic is not None else False
    formatting['underline'] = run.underline if run.underline is not None else False
    formatting['font_size'] = run.font.size
    formatting['font_name'] = run.font.name
    return formatting

# Apply formatting to a target run
def apply_formatting(run, formatting):
    run.bold = formatting.get('bold', False)
    run.italic = formatting.get('italic', False)
    run.underline = formatting.get('underline', False)
    run.font.size = formatting.get('font_size')
    run.font.name = formatting.get('font_name')

# Check if the paragraph is a list paragraph
def is_list_paragraph(paragraph):
    p = paragraph._element
    pPr = p.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        return numPr is not None
    return False

# Apply numbering list formatting (font and size) to the paragraph's list numbers
def apply_list_number_formatting(paragraph, font_name, font_size):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size

# Create a custom numbering style for lists in the target document
def create_custom_list_style(target_doc, style_name='CustomListNumber', font_name='Comic Sans MS', font_size=Pt(12)):
    styles = target_doc.styles

    # Check if the style already exists, return if it does
    if style_name in [s.name for s in styles]:
        return styles[style_name]

    # Create a new numbering style
    new_style = styles.add_style(style_name, 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
    new_style.font.size = font_size  # Set custom size
    new_style.font.name = font_name  # Set custom font
    return new_style

# Apply list formatting from source to target
def apply_list_formatting(source_paragraph, target_paragraph, target_doc):
    font_name = source_paragraph.runs[0].font.name
    font_size = source_paragraph.runs[0].font.size

    # Create and apply custom list style
    custom_list_style = create_custom_list_style(target_doc, font_name=font_name, font_size=font_size)
    target_paragraph.style = custom_list_style

    # Apply number formatting to list
    apply_list_number_formatting(target_paragraph, font_name, font_size)

# Apply paragraph formatting (non-list) from source to target
def apply_paragraph_formatting(source_paragraph, target_paragraph):
    target_paragraph.text = source_paragraph.text

    for source_run, target_run in zip(source_paragraph.runs, target_paragraph.runs):
        formatting = extract_formatting(source_run)
        apply_formatting(target_run, formatting)

# Detect and apply formatting based on text type (list, body, etc.)
def apply_text_type_formatting(source_doc, target_doc):
    for source_paragraph in source_doc.paragraphs:
        # Find a paragraph in target doc that matches the source's paragraph type
        for target_paragraph in target_doc.paragraphs:
            if is_list_paragraph(source_paragraph):
                # If source paragraph is a list, apply list formatting
                apply_list_formatting(source_paragraph, target_paragraph, target_doc)
            else:
                # Apply normal paragraph formatting for body text
                apply_paragraph_formatting(source_paragraph, target_paragraph)

# Apply the formatting
apply_text_type_formatting(source_doc, target_doc)

# Save the target document
target_doc.save('target.docx')
