from docx import Document
from docx.shared import Pt # for font size
from docx.oxml import OxmlElement # add page number field
from docx.oxml.ns import qn  # Import for XML namespaces

# extract formatting from paragraph
def extract_formatting(run):
    formatting = {}
    formatting['bold'] = run.bold if run.bold is not None else False
    formatting['italic'] = run.italic if run.italic is not None else False
    formatting['underline'] = run.underline if run.underline is not None else False
    formatting['font_size'] = run.font.size
    formatting['font_name'] = run.font.name
    formatting['font'] = run.font
    # print('formatting: ', formatting)
    return formatting

# apply formatting to paragraph
def apply_formatting(run, formatting):
    run.bold = formatting.get('bold', False)
    run.italic = formatting.get('italic', False)
    run.underline = formatting.get('underline', False)
    run.font.size = formatting.get('font_size')
    run.font.name = formatting.get('font_name')

# check if is list
def is_list_paragraph(paragraph):
    # the paragraph element is w:p
    p = paragraph._element
    # need to find the subset w:pPr, which contains numPr
    pPr = p.find(qn('w:pPr'))
    # w:numPr is a subset within w:pPr, so check
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        return numPr is not None
    return False

# extract formatting of list
def extract_list_formatting(paragraph):
    list_formatting = {}
    # check if it's a list
    if is_list_paragraph(paragraph):
        print('extracting list formatting')
        print('paragraph: ', paragraph.runs[0].font.name)

        p = paragraph._element
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                # Extract numbering level (ilvl) and numId (list ID)
                ilvl = numPr.find(qn('w:ilvl'))
                numId = numPr.find(qn('w:numId'))
                if ilvl is not None:
                    list_formatting['ilvl'] = ilvl.get(qn('w:val'))
                if numId is not None:
                    list_formatting['numId'] = numId.get(qn('w:val'))
    for run in paragraph.runs:
        list_formatting['font_size'] = run.font.size
        list_formatting['font_name'] = run.font.name
    return list_formatting

# apply list formatting
def apply_list_formatting(paragraph, list_formatting):
    if list_formatting:
        # print('list formatting: ', list)

        # Create the numbering properties in the target paragraph
        pPr = paragraph._element.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        
        # Apply numId and ilvl (indentation level)
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), list_formatting['numId'])
        
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), list_formatting['ilvl'])

        print('numId: ', numId)
        print('ilvl: ', ilvl)

        numPr.append(numId)
        numPr.append(ilvl)
        pPr.append(numPr)

        for run in paragraph.runs:
            run.font.size = list_formatting.get('font_size')
            run.font.name = list_formatting.get('font_name')

# insert page number to match source doc
def insert_page_number(paragraph):
    run = paragraph.add_run() 

    # create field for page number 
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1) # add to the run

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2) 

# check if page number exists
def contains_page_number(paragraph):
    for run in paragraph.runs:
        for element in run._r:
            if element.tag.endswith('instrText') and 'PAGE' in element.text:
                return True
    return False

# extract header
def extract_header(document):
    # Ensure the document has sections
    if len(document.sections) > 0:
        return document.sections[0].header
    else:
        raise ValueError("The document has no sections.")

# extract footer
def extract_footer(document):
    # Ensure the document has sections
    if len(document.sections) > 0:
        return document.sections[0].footer
    else:
        raise ValueError("The document has no sections.")
    
# Clear and apply header to the target document
def apply_header(target_document, source_header):
    target_header = target_document.sections[0].header
    # if target header exists, remove existing paragraph
    if target_header:
        # Clear existing header paragraphs
        for paragraph in target_header.paragraphs:
            p = paragraph._element  # Access underlying XML element
            p.getparent().remove(p)
    # Copy content from the source header
    for paragraph in source_header.paragraphs:
        new_paragraph = target_header.add_paragraph(paragraph.text)
        # ensure the alignment of the target footer aligns with source footer
        new_paragraph.alignment = paragraph.alignment
        for source_run, new_run in zip(paragraph.runs, new_paragraph.runs):
            formatting = extract_formatting(source_run)
            apply_formatting(new_run, formatting)
        # Check if the source footer contains the page number and add it to the target
        if contains_page_number(paragraph):
            insert_page_number(new_paragraph)
            print('yes')

# Clear and apply footer to the target document
def apply_footer(target_document, source_footer):
    target_footer = target_document.sections[0].footer
    # if target footer exists, remove existing paragraph
    if target_footer:
        # Clear existing footer paragraphs
        for paragraph in target_footer.paragraphs:
            p = paragraph._element  # Access underlying XML element
            p.getparent().remove(p)
    # Copy content from the source footer
    for paragraph in source_footer.paragraphs:
        new_paragraph = target_footer.add_paragraph(paragraph.text)
        # ensure the alignment of the target footer aligns with source footer
        new_paragraph.alignment = paragraph.alignment
        for source_run, new_run in zip(paragraph.runs, new_paragraph.runs):
            formatting = extract_formatting(source_run)
            apply_formatting(new_run, formatting)
        # Check if the source footer contains the page number and add it to the target
        if contains_page_number(paragraph):
            insert_page_number(new_paragraph)
            print('yes')

# apply formatting for paragraphs, including list formatting
def apply_paragraph_formatting(source_paragraph, target_paragraph):

    # check if source paragraph is list
    if is_list_paragraph(source_paragraph):
        # extract and apply list formattings
        list_formatting = extract_list_formatting(source_paragraph)
        apply_list_formatting(target_paragraph, list_formatting)

    # ensure the alignment of the target footer aligns with source footer
    target_paragraph.alignment = source_paragraph.alignment
    
    # extract and apply normal formatting to normal text
    if source_paragraph.runs:
        for source_run, target_run in zip(source_paragraph.runs, target_paragraph.runs):
            formatting = extract_formatting(source_run)
            apply_formatting(target_run, formatting)

# load source and target documents
source_doc = Document('source.docx')
target_doc = Document('target.docx')

# Iterate through paragraphs and print XML for list paragraphs
for i, paragraph in enumerate(source_doc.paragraphs):
    # print(paragraph._element.xml)
    if is_list_paragraph(paragraph):
        print(f"Paragraph {i} is part of a list.")
        print(paragraph._element.xml)
    else:
        print(f"Paragraph {i} is NOT part of a list.")

# Ensure both source and target documents have at least one section
if len(source_doc.sections) == 0:
    raise ValueError("The source document has no sections.")
if len(target_doc.sections) == 0:
    raise ValueError("The target document has no sections.")

# Formatting:
for source_paragraph, target_paragraph in zip(source_doc.paragraphs, target_doc.paragraphs):
    apply_paragraph_formatting(source_paragraph, target_paragraph)

# Header and footer:

# extract header and footer from source document
header = extract_header(source_doc)
footer = extract_footer(source_doc)

# apply header and footer to target document
apply_header(target_doc, header)
apply_footer(target_doc, footer)

# save target document
target_doc.save('target.docx')